from docx import Document
from docx.exceptions import PasswordRequiredError

class DocReader:
    def __init__(self, docx_file_path):
        self.docx_file_path = docx_file_path

    def is_password_protected(self):
        try:
            # Attempt to open the document without specifying a password
            Document(self.docx_file_path)
            return False  # Document opened successfully, not password-protected
        except PasswordRequiredError:
            return True  # Document is password-protected

    def read_document(self):
        if self.is_password_protected():
            return None  # Return None if the document is password-protected

        try:
            doc = Document(self.docx_file_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            return text
        except Exception as e:
            return str(e)  # Handle any exceptions during document reading
