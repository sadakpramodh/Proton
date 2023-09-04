'''
Read the pdf, csv, word and excel files!
'''

import fitz
import PyPDF2
import os
import pytesseract
from pdf2image import convert_from_path

from docx import Document
from docx.exceptions import PasswordRequiredError

import pandas as pd



class PDFReader:
    def __init__(self, file_path):
        self.file_path = file_path

    def check_pdftype(file_path):
        # check weather pdf is OCR based or text based
        doc = fitz.open(file_path)
        for page in doc:
            if len(page.get_text()) == 0:
                print("This page is an image")
            else:
                print("This page contains text")

    def check_password(file_path, password):
        with open(file_path, 'rb') as f:
            pdf = PyPDF2.PdfFileReader(f)
            if pdf.isEncrypted:
                try:
                    pdf.decrypt(password)
                    return True
                except PyPDF2.utils.PdfReadError:
                    return False
            else:
                return False
            
    # Function to read PDF that are in image format and convert to text
    def read_pdf_image(pdf_path):
        images = convert_from_path(pdf_path)
        text = ""
        for page in images:
            text += pytesseract.image_to_string(page)
        return text
    # Function to read PDF and convert to text
    def read_pdf_text(pdf_path):
        doc = fitz.open(pdf_path)
        text = ""
        for page in doc:
            text += page.get_text()
        return text
 

class DocReader:
    def __init__(self, docx_file_path):
        self.docx_file_path = docx_file_path

    def is_password_protected(self):
        try:
            # Attempt to open the document without specifying a password
            Document(self.docx_file_path)
            return False  # Document opened successfully, not password-protected
        except PasswordRequiredError:
            return True  # Document is password-protected

    def read_document(self):
        if self.is_password_protected():
            return None  # Return None if the document is password-protected

        try:
            doc = Document(self.docx_file_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            return text
        except Exception as e:
            return str(e)  # Handle any exceptions during document reading


class ExcelReader:
    def __init__(self, excel_file_path):
        self.excel_file_path = excel_file_path

    def is_password_protected(self):
        try:
            # Attempt to read the Excel file
            pd.read_excel(self.excel_file_path)
            return False  # Excel file opened successfully, not password-protected
        except Exception as e:
            error_message = str(e)
            if "password" in error_message.lower():
                return True  # Excel file is password-protected
            else:
                return False  # Other error occurred, not password-protected

    def read_excel(self):
        if self.is_password_protected():
            return None  # Return None if the Excel file is password-protected

        try:
            df = pd.read_excel(self.excel_file_path)
            return df
        except Exception as e:
            return str(e)  # Handle any exceptions during Excel file reading

class CSVReader:
    def __init__(self, csv_file_path):
        self.csv_file_path = csv_file_path

    def read_csv(self):
        try:
            df = pd.read_csv(self.csv_file_path)
            return df
        except Exception as e:
            return str(e)  # Handle any exceptions during CSV file reading
